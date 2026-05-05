# ============================================================
# 22_supply_courses.py
# 用途：V5 — 课程粒度四元组抽取（Option C 超细粒度）
#       Phase α：从 docx 解析每门课的"课程名 / 学时 / 课程简介"
#       Phase β：用 LLM 对每门课直接抽取四元组（工具/技能/能力/知识点）
# 输出：data/v5_supply_courses_raw.parquet（α 输出）
#       data/v5_supply_courses.parquet     （β 输出）
# 依赖：scripts/llm_client.py 的底层 _call_llm / _parse_json_from_response
#       scripts/15_v4_supply_extraction.py 的 MAJORS 清单
#       python-docx
# ============================================================

import os
import sys
import re
import math
import pathlib
import json
import argparse
from typing import List, Dict, Optional

import pandas as pd
import numpy as np
import docx as docx_lib

sys.path.insert(0, ".")
import importlib.util
spec = importlib.util.spec_from_file_location(
    "m15", "scripts/15_v4_supply_extraction.py"
)
m15 = importlib.util.module_from_spec(spec)
spec.loader.exec_module(m15)
MAJORS = m15.MAJORS
PLAN_ROOT = pathlib.Path("源数据/人才培养方案")

from scripts.llm_client import (
    _get_client, _call_llm, _parse_json_from_response,
    _save_checkpoint, _load_checkpoint, DEFAULT_MODEL,
)

OUT_RAW = pathlib.Path("data/v5_supply_courses_raw.parquet")
OUT_FINAL = pathlib.Path("data/v5_supply_courses.parquet")
OUT_RAW.parent.mkdir(parents=True, exist_ok=True)


# ============================================================
# Phase α — 从 docx 中识别课程清单
# ============================================================
COURSE_NAME_KEYWORDS = (
    "程序设计", "原理", "技术", "工程", "系统", "管理", "实训", "实习", "实践",
    "应用", "导论", "基础", "概论", "训练", "实验", "设计", "分析",
    "建模", "操作", "开发", "算法", "网络", "数据", "云", "通信", "电路",
    "机器", "工业", "智能", "嵌入", "服务", "运维", "测试", "项目", "毕业",
    "建筑", "环境", "园林", "财务", "会计", "金融", "营销", "电商", "市场",
    "外语", "英语", "教育", "护理", "体育", "传媒", "媒体", "动画", "影视",
    "结构", "材料", "理论", "方法", "学", "组织", "战略", "运营", "国际",
)
NON_COURSE_TOKENS = (
    "学时", "学分", "考核", "总学时", "周学时",
    "实践学时", "理论学时", "学期", "授课方式", "开课", "成绩",
    "教学方法", "教学目标", "课程目标", "教学评价", "评价方式",
    "课程性质", "课程类别",
)


def parse_docx_blocks(path: pathlib.Path) -> List:
    """解析 docx 返回 [(类型, 内容), ...]，类型 ∈ {'para','table_row'}。"""
    if not path.exists():
        return []
    try:
        doc = docx_lib.Document(str(path))
    except Exception:
        return []
    blocks = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            blocks.append(("para", t))
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                blocks.append(("table_row", cells))
    return blocks


def detect_course_from_row(cells: List[str]) -> Optional[Dict]:
    """严格识别一门课程：必须同时具备
       (1) 课程名（4-25 字、含学科关键词、不在排除词单中）
       (2) 学时（数字 16-320 之间，是高职典型课程学时范围）
    """
    s = " | ".join(cells)
    if any(t in s for t in NON_COURSE_TOKENS) and not any(k in s for k in COURSE_NAME_KEYWORDS):
        return None
    course_name = None
    for c in cells:
        c = c.strip()
        # 严格的课程名规则：4—25 字，含关键词，不含括号注释里的元信息
        if 4 <= len(c) <= 25 and any(k in c for k in COURSE_NAME_KEYWORDS):
            # 排除像"开设学期"、"理论学时"等表头/说明
            if any(x in c for x in ["开设", "总学", "授课", "分配", "任课"]):
                continue
            course_name = c
            break
    if not course_name:
        return None
    # 学时必须在 16—320 范围（高职典型课程学时；过小/过大都说明不是课程行）
    hours = None
    for c in cells:
        c2 = c.strip().replace(".", "")
        if c2.isdigit():
            v = int(c2)
            if 16 <= v <= 320:
                hours = v
                break
    if hours is None:
        return None  # 没有学时的"课程"通常是矩阵单元，过滤
    return {"course_name": course_name, "hours": hours, "raw_row": s}


def extract_course_descriptions(blocks: List) -> Dict[str, str]:
    """尝试匹配课程名与简介（取同一行最长的那格）。"""
    desc = {}
    for kind, payload in blocks:
        if kind != "table_row":
            continue
        cells = payload
        if len(cells) < 2:
            continue
        name = None
        for c in cells:
            if 3 <= len(c.strip()) <= 30 and any(k in c for k in COURSE_NAME_KEYWORDS):
                name = c.strip()
                break
        if not name:
            continue
        longest = max(cells, key=lambda x: len(x.strip()))
        if len(longest.strip()) > 50 and longest.strip() != name:
            desc[name] = longest.strip()[:1500]
    return desc


def run_phase_alpha():
    print("=" * 60)
    print("V5 Step 22 · Phase α — 课程清单解析")
    print("=" * 60)

    rows = []
    for major, college, level, main, subs in MAJORS:
        files = ([main] if main else []) + (subs or [])
        all_blocks = []
        for f in files:
            full = PLAN_ROOT / f
            all_blocks.extend(parse_docx_blocks(full))

        seen, courses = set(), []
        for kind, payload in all_blocks:
            if kind == "table_row":
                c = detect_course_from_row(payload)
                if c and c["course_name"] not in seen:
                    seen.add(c["course_name"])
                    courses.append(c)

        descs = extract_course_descriptions(all_blocks)

        for c in courses:
            rows.append({
                "major": major, "college": college, "level": level,
                "course_name": c["course_name"], "hours": c["hours"],
                "description": descs.get(c["course_name"], ""),
                "raw_row": c["raw_row"],
            })
        n_with_desc = sum(1 for c in courses if descs.get(c["course_name"]))
        print(f"  {major:25s}（{college}）：{len(courses):3d} 门课，简介 {n_with_desc:3d} 门")

    df = pd.DataFrame(rows)
    df = df.drop_duplicates(subset=["major", "course_name"]).reset_index(drop=True)
    df["course_id"] = [f"C_{i:04d}" for i in range(len(df))]
    df.to_parquet(OUT_RAW, index=False)
    print(f"\nα 完成：{len(df)} 门唯一课程（54 专业）→ {OUT_RAW}")
    return df


# ============================================================
# Phase β — LLM 抽取四元组
# ============================================================
SYSTEM_QUAD = (
    "你是专业的高职课程分析师。请从课程信息中抽取四个维度的关键要素："
    "（1）工具/软件/平台（具体产品名）；（2）技能（可操作的具体技能）；"
    "（3）能力（认知/职业能力，如分析、协作、问题解决）；"
    "（4）知识点（学科理论概念、原理）。"
    "只返回JSON，不要任何额外说明。"
)
USER_QUAD_TPL = (
    "对以下{count}门课程的'课程名+课程简介'文本，分别抽取四元组（工具/技能/能力/知识点）。\n\n"
    "工具：编程语言、设计/工程软件、行业平台等具体产品（Python、CAD、PLC、ERP、Photoshop、SolidWorks 等）。\n"
    "技能：可操作的具体技能（电路调试、数据库查询、机械装配、需求分析、报表编制等）。\n"
    "能力：认知或职业能力（分析能力、沟通协调、问题解决、创新思维、团队协作等）。\n"
    "知识点：学科理论或概念（离散数学、机器学习、电磁学、营销原理、会计准则等）。\n\n"
    "如果某维度无内容，对应字段输出 []。\n\n"
    "输出格式（JSON 数组）：\n"
    '[{{"id": "课程ID", "tools": [...], "skills": [...], "abilities": [...], "knowledge": [...]}}, ...]\n\n'
    "课程信息：\n{texts}"
)


def make_text(row) -> str:
    if row["description"]:
        return f"【课程名】{row['course_name']}\n【主要内容】{row['description']}"
    return f"【课程名】{row['course_name']}\n【说明】（无明示课程简介，仅根据课程名推断）"


def run_phase_beta(subset: int = 0, batch_size: int = 5):
    print("=" * 60)
    print("V5 Step 22 · Phase β — LLM 四元组抽取")
    print("=" * 60)

    df = pd.read_parquet(OUT_RAW)
    if subset > 0:
        df = df.head(subset).copy()
        print(f"  ⚠️ 测试模式：仅处理前 {subset} 门课")
    df["分析文本"] = df.apply(make_text, axis=1)

    n = len(df)
    print(f"  待抽取课程数：{n}（含简介 {(df['description'].str.len() > 0).sum()} 门）")

    task = "v5_courses_quad"
    ckpt_dir = pathlib.Path("data/v2_checkpoints")
    # 清除旧检查点（reset）
    for p in ckpt_dir.glob(f"{task}_batch*.jsonl"):
        p.unlink()

    client = _get_client()
    records = df[["course_id", "分析文本"]].to_dict("records")
    total = math.ceil(n / batch_size)

    all_results = []
    for bi in range(total):
        batch = records[bi*batch_size:(bi+1)*batch_size]
        texts_block = "\n\n".join(
            [f'[{r["course_id"]}] {str(r["分析文本"])[:1500]}' for r in batch]
        )
        user = USER_QUAD_TPL.format(count=len(batch), texts=texts_block)
        msgs = [
            {"role": "system", "content": SYSTEM_QUAD},
            {"role": "user", "content": user},
        ]
        try:
            raw = _call_llm(client, msgs, DEFAULT_MODEL)
            parsed = _parse_json_from_response(raw)
            batch_results = []
            if isinstance(parsed, list):
                for item in parsed:
                    if isinstance(item, dict):
                        batch_results.append({
                            "course_id": item.get("id") or item.get("course_id"),
                            "tools": list(item.get("tools", []) or []),
                            "skills": list(item.get("skills", []) or []),
                            "abilities": list(item.get("abilities", []) or []),
                            "knowledge": list(item.get("knowledge", []) or []),
                        })
            if not batch_results:
                batch_results = [
                    {"course_id": r["course_id"], "tools": [], "skills": [], "abilities": [], "knowledge": []}
                    for r in batch
                ]
        except Exception as e:
            print(f"    第 {bi} 批失败：{e}")
            batch_results = [
                {"course_id": r["course_id"], "tools": [], "skills": [], "abilities": [], "knowledge": []}
                for r in batch
            ]
        all_results.extend(batch_results)
        _save_checkpoint(task, bi, batch_results)
        if (bi + 1) % 20 == 0 or (bi + 1) == total:
            print(f"    进度 {bi+1}/{total} ({(bi+1)/total:.1%})")

    quad_df = pd.DataFrame(all_results)
    out = df[["course_id", "major", "college", "level", "course_name", "hours", "description"]].merge(
        quad_df, on="course_id", how="left"
    )
    for col in ["tools", "skills", "abilities", "knowledge"]:
        out[col] = out[col].apply(
            lambda x: list(x) if hasattr(x, "__iter__") and not isinstance(x, str) else []
        )

    out.to_parquet(OUT_FINAL, index=False)
    print(f"\nβ 完成 → {OUT_FINAL}")
    for col in ["tools", "skills", "abilities", "knowledge"]:
        n_have = sum(1 for x in out[col] if x)
        print(f"  {col} 非空：{n_have}/{len(out)} ({n_have/len(out)*100:.1f}%)")
    return out


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--phase", choices=["a", "b", "all"], default="all")
    parser.add_argument("--subset", type=int, default=0)
    parser.add_argument("--batch_size", type=int, default=5)
    args = parser.parse_args()
    if args.phase in ("a", "all"):
        run_phase_alpha()
    if args.phase in ("b", "all"):
        run_phase_beta(subset=args.subset, batch_size=args.batch_size)
    print("\n=== Step 22 完成 ===")
