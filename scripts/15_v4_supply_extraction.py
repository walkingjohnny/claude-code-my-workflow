# ============================================================
# 15_supply_side_extraction.py
# 用途：V3 研究 — 供给侧（55 个专业的人才培养方案）LLM 文本挖掘
#       为保证与需求侧（智联招聘 16,501 条记录）可对接，
#       本脚本使用与 09/10/12 完全一致的 prompt 与归一化映射。
# 输出：data/v4_supply_majors_raw.parquet（解析后的聚合文本）
#       data/v4_supply_majors.parquet（LLM 抽取后的最终供给侧画像）
# 依赖：scripts/llm_client.py, scripts/_common.py
#       源数据/人才培养方案/{职业本科, 高职专科}/...
# ============================================================

import os
import sys
import pathlib
import re
import json
from typing import Optional

import pandas as pd
import numpy as np

# 项目工具
sys.path.insert(0, ".")
from scripts.llm_client import batch_extract, batch_rate

# 文档解析
import docx as docx_lib  # python-docx

# --- 常量 ---
PLAN_ROOT = pathlib.Path("源数据/人才培养方案")
OUT_RAW = pathlib.Path("data/v4_supply_majors_raw.parquet")
OUT_FINAL = pathlib.Path("data/v4_supply_majors.parquet")
OUT_RAW.parent.mkdir(parents=True, exist_ok=True)


# ============================================================
# 1. 归并映射（55 个独立专业）
#    — 用户已确认的清单
#    每条 = (专业名, 学院, 教育层次, 主文件相对路径或None, [子轨道文件相对路径])
# ============================================================
MAJORS = [
    # ===== 职业本科（5 个，无子轨道）=====
    ("嵌入式技术", "中德机器人学院", "职业本科",
     "职业本科/中德机器人学院：2025级嵌入式技术本科专业人才培养方案.docx", []),
    ("现代通信工程", "信息与通信学院", "职业本科",
     "职业本科/信息与通信学院：2025级现代通信工程本科专业人才培养方案.docx", []),
    ("集成电路工程技术", "微电子学院", "职业本科",
     "职业本科/微电子学院：2025级集成电路工程技术本科专业人才培养方案.docx", []),
    ("智能制造工程技术", "智能制造与装备学院", "职业本科",
     "职业本科/智能制造与装备学院：2025级智能制造工程技术专业人才培养方案.docx", []),
    ("软件工程技术", "计算机与软件学院", "职业本科",
     "职业本科/计算机与软件学院：2025级软件工程技术本科专业人才培养方案.docx", []),

    # ===== 1. 计算机与软件学院（7 个）=====
    ("信息安全技术应用", "计算机与软件学院", "高职专科",
     "高职专科/1.计算机与软件学院/计算机与软件学院：信息安全技术应用专业人才培养方案.docx", []),
    ("区块链技术应用", "计算机与软件学院", "高职专科",
     "高职专科/1.计算机与软件学院/计算机与软件学院：区块链技术应用专业人才培养方案.docx", []),
    ("移动互联应用技术", "计算机与软件学院", "高职专科",
     None,
     ["高职专科/1.计算机与软件学院/计算机与软件学院：移动互联应用技术专业（工业互联应用开发方向）人才培养方案.docx",
      "高职专科/1.计算机与软件学院/计算机与软件学院：移动互联应用技术专业（移动互联应用开发方向）人才培养方案.docx"]),
    ("计算机应用技术", "计算机与软件学院", "高职专科",
     "高职专科/1.计算机与软件学院/计算机与软件学院：计算机应用技术专业人才培养方案0624.docx",
     ["高职专科/1.计算机与软件学院/计算机与软件学院：计算机应用技术专业（奇点班）人才培养方案.docx"]),
    ("计算机网络技术", "计算机与软件学院", "高职专科",
     "高职专科/1.计算机与软件学院/计算机与软件学院：计算机网络技术专业人才培养方案0624.docx", []),
    ("软件技术", "计算机与软件学院", "高职专科",
     None,
     ["高职专科/1.计算机与软件学院/计算机与软件学院：软件技术专业（Java技术开发方向）人才培养方案.docx",
      "高职专科/1.计算机与软件学院/计算机与软件学院：软件技术专业（Python技术开发方向）人才培养方案.docx",
      "高职专科/1.计算机与软件学院/计算机与软件学院：软件技术专业（工业软件应用开发方向）人才培养方案.docx",
      "高职专科/1.计算机与软件学院/计算机与软件学院：软件技术专业（腾飞班）人才培养方案.docx"]),

    # ===== 2. 信息与通信学院（4 个）=====
    ("智能互联网络技术", "信息与通信学院", "高职专科",
     "高职专科/2.信息与通信学院/信息与通信学院：智能互联网络技术专业培养方案.docx",
     ["高职专科/2.信息与通信学院/信息与通信学院：智能互联网络技术专业（鲲鹏班）培养方案.docx"]),
    ("汽车智能技术", "信息与通信学院", "高职专科",
     "高职专科/2.信息与通信学院/信息与通信学院：汽车智能技术专业培养方案.docx",
     ["高职专科/2.信息与通信学院/信息与通信学院：汽车智能技术专业（鸿蒙班）培养方案.docx"]),
    ("物联网应用技术", "信息与通信学院", "高职专科",
     "高职专科/2.信息与通信学院/信息与通信学院：物联网应用技术专业培养方案.docx",
     ["高职专科/2.信息与通信学院/信息与通信学院：物联网应用技术专业（鸿蒙班）培养方案.docx"]),
    ("现代移动通信技术", "信息与通信学院", "高职专科",
     None,
     ["高职专科/2.信息与通信学院/信息与通信学院：现代移动通信技术专业（5G智能网络及应用方向）培养方案.docx",
      "高职专科/2.信息与通信学院/信息与通信学院：现代移动通信技术专业（移动智能终端及应用方向）培养方案.docx"]),

    # ===== 3. 微电子学院（3 个）=====
    ("智能产品开发与应用", "微电子学院", "高职专科",
     "高职专科/3.微电子学院/微电子学院：智能产品开发与应用专业培养方案.docx", []),
    ("智能光电技术应用", "微电子学院", "高职专科",
     "高职专科/3.微电子学院/微电子学院：智能光电技术应用专业培养方案.docx", []),
    ("集成电路技术", "微电子学院", "高职专科",
     "高职专科/3.微电子学院/微电子学院：集成电路技术专业培养方案.docx",
     ["高职专科/3.微电子学院/微电子学院：集成电路技术专业（特色班）培养方案.docx"]),

    # ===== 4. 人工智能学院（3 个）=====
    ("云计算技术应用", "人工智能学院", "高职专科",
     "高职专科/4.人工智能学院/人工智能学院：云计算技术应用专业培养方案.docx",
     ["高职专科/4.人工智能学院/人工智能学院：云计算技术应用专业（云创班）培养方案.docx"]),
    ("人工智能技术应用", "人工智能学院", "高职专科",
     "高职专科/4.人工智能学院/人工智能学院：人工智能技术应用专业培养方案.docx",
     ["高职专科/4.人工智能学院/人工智能学院：人工智能技术应用专业（智创班）培养方案.docx"]),
    ("大数据技术", "人工智能学院", "高职专科",
     "高职专科/4.人工智能学院/人工智能学院：大数据技术专业培养方案.docx",
     ["高职专科/4.人工智能学院/人工智能学院：大数据技术专业（数创班）培养方案.docx"]),

    # ===== 5. 数字媒体学院（5 个）=====
    ("广播影视节目制作", "数字媒体学院", "高职专科",
     "高职专科/5.数字媒体学院/数字媒体学院：广播影视节目制作.docx", []),
    ("数字媒体技术", "数字媒体学院", "高职专科",
     "高职专科/5.数字媒体学院/数字媒体学院：数字媒体技术.docx",
     ["高职专科/5.数字媒体学院/数字媒体学院：数字媒体技术（易数班）.docx"]),
    ("数字媒体艺术设计", "数字媒体学院", "高职专科",
     None,
     ["高职专科/5.数字媒体学院/数字媒体学院：数字媒体艺术设计（产品艺术设计方向）.docx",
      "高职专科/5.数字媒体学院/数字媒体学院：数字媒体艺术设计（视觉传达设计方向）.docx"]),
    ("环境艺术设计", "数字媒体学院", "高职专科",
     "高职专科/5.数字媒体学院/数字媒体学院：环境艺术设计.docx", []),
    ("虚拟现实技术应用", "数字媒体学院", "高职专科",
     "高职专科/5.数字媒体学院/数字媒体学院：虚拟现实技术应用.docx", []),

    # ===== 6. 智能制造与装备学院（5 个）=====
    ("工业设计", "智能制造与装备学院", "高职专科",
     "高职专科/6.智能制造与装备学院/智能制造与装备学院：工业设计专业培养方案.docx", []),
    ("工业软件开发技术", "智能制造与装备学院", "高职专科",
     "高职专科/6.智能制造与装备学院/智能制造与装备学院：工业软件开发技术专业培养方案.docx",
     ["高职专科/6.智能制造与装备学院/智能制造与装备学院：工业软件开发技术专业（汇信班）培养方案.docx"]),
    ("智能光电制造技术", "智能制造与装备学院", "高职专科",
     "高职专科/6.智能制造与装备学院/智能制造与装备学院：智能光电制造技术专业培养方案.docx", []),
    ("智能控制技术", "智能制造与装备学院", "高职专科",
     "高职专科/6.智能制造与装备学院/智能制造与装备学院：智能控制技术专业培养方案.docx",
     ["高职专科/6.智能制造与装备学院/智能制造与装备学院：智能控制技术专业（汇信班）培养方案.docx"]),
    ("机械设计与制造", "智能制造与装备学院", "高职专科",
     "高职专科/6.智能制造与装备学院/智能制造与装备学院：机械设计与制造专业培养方案.docx", []),

    # ===== 7. 交通与环境学院（5 个）=====
    ("园林工程技术", "交通与环境学院", "高职专科",
     "高职专科/7.交通与环境学院/交通与环境学院：园林工程技术专业培养方案.docx", []),
    ("城市轨道交通运营管理", "交通与环境学院", "高职专科",
     "高职专科/7.交通与环境学院/交通与环境学院：城市轨道交通运营管理专业培养方案.docx", []),
    ("智能建造技术", "交通与环境学院", "高职专科",
     "高职专科/7.交通与环境学院/交通与环境学院：智能建造技术专业培养方案.docx", []),
    ("环境工程技术", "交通与环境学院", "高职专科",
     "高职专科/7.交通与环境学院/交通与环境学院：环境工程技术专业培养方案.docx",
     ["高职专科/7.交通与环境学院/交通与环境学院：环境工程技术专业（求实班）培养方案.docx"]),
    ("环境监测技术", "交通与环境学院", "高职专科",
     "高职专科/7.交通与环境学院/交通与环境学院：环境监测技术专业培养方案.docx", []),

    # ===== 8. 管理学院（6 个）=====
    ("关务与外贸服务", "管理学院", "高职专科",
     "高职专科/8.管理学院/管理学院：关务与外贸服务专业培养方案.docx", []),
    ("国际商务", "管理学院", "高职专科",
     "高职专科/8.管理学院/管理学院：国际商务专业培养方案.docx", []),
    ("工商企业管理", "管理学院", "高职专科",
     "高职专科/8.管理学院/管理学院：工商企业管理专业培养方案.docx", []),
    ("文化产业经营与管理", "管理学院", "高职专科",
     "高职专科/8.管理学院/管理学院：文化产业经营与管理专业培养方案.docx", []),
    ("现代物流管理", "管理学院", "高职专科",
     "高职专科/8.管理学院/管理学院：现代物流管理专业培养方案.docx",
     ["高职专科/8.管理学院/管理学院：现代物流管理专业（智慧物流工程特色班）培养方案.docx"]),
    ("电子商务", "管理学院", "高职专科",
     "高职专科/8.管理学院/管理学院：电子商务专业培养方案.docx", []),

    # ===== 9. 财经学院（4 个）=====
    ("大数据与会计", "财经学院", "高职专科",
     "高职专科/9.财经学院/财经学院：大数据与会计专业人才培养方案.docx",
     ["高职专科/9.财经学院/财经学院：大数据与会计专业（正保数智特色班）人才培养方案.docx"]),
    ("财富管理", "财经学院", "高职专科",
     "高职专科/9.财经学院/财经学院：财富管理专业培养方案.docx",
     ["高职专科/9.财经学院/财经学院：财富管理专业（圳金班）培养方案.docx"]),
    ("财税大数据应用", "财经学院", "高职专科",
     "高职专科/9.财经学院/财经学院：财税大数据应用专业培养方案.docx", []),
    ("金融服务与管理", "财经学院", "高职专科",
     "高职专科/9.财经学院/财经学院：金融服务与管理专业培养方案.docx",
     ["高职专科/9.财经学院/财经学院：金融服务与管理专业（数智金融班）培养方案.docx"]),

    # ===== 10. 应用外语学院（3 个）=====
    ("商务英语", "应用外语学院", "高职专科",
     "高职专科/10.应用外语学院/应用外语学院：商务英语专业培养方案.docx",
     ["高职专科/10.应用外语学院/应用外语学院：商务英语专业（信言涉外法律特色班）培养方案.docx"]),
    ("学前教育", "应用外语学院", "高职专科",
     "高职专科/10.应用外语学院/应用外语学院：学前教育专业人才培养方案.docx",
     ["高职专科/10.应用外语学院/应用外语学院：学前教育专业（深业智慧照护班）培养方案.docx"]),
    ("现代文秘", "应用外语学院", "高职专科",
     "高职专科/10.应用外语学院/应用外语学院：现代文秘专业人才培养方案.docx", []),

    # ===== 11. 中德机器人学院（4 个）=====
    ("工业互联网技术", "中德机器人学院", "高职专科",
     "高职专科/11.中德机器人学院/中德机器人学院：工业互联网技术专业培养方案.docx", []),
    ("工业机器人技术", "中德机器人学院", "高职专科",
     "高职专科/11.中德机器人学院/中德机器人学院：工业机器人技术专业培养方案.docx",
     ["高职专科/11.中德机器人学院/中德机器人学院：工业机器人技术专业（汇信班）培养方案.docx"]),
    ("无人机应用技术", "中德机器人学院", "高职专科",
     "高职专科/11.中德机器人学院/中德机器人学院：无人机应用技术专业培养方案.docx", []),
    ("智能机器人技术", "中德机器人学院", "高职专科",
     "高职专科/11.中德机器人学院/中德机器人学院：智能机器人技术专业培养方案.docx", []),

    # ===== 12. 体育运动学院（1 个）=====
    ("体育艺术表演", "体育运动学院", "高职专科",
     "高职专科/12.体育运动学院/体育运动学院：体育艺术表演专业培养方案.docx", []),
]


# ============================================================
# 2. DOCX 解析
# ============================================================
def parse_docx(path: pathlib.Path) -> str:
    """解析 docx 提取所有段落+表格文本，返回单个长字符串。"""
    if not path.exists():
        return ""
    try:
        doc = docx_lib.Document(str(path))
    except Exception as e:
        print(f"  ⚠ 解析失败 {path.name}: {e}")
        return ""
    chunks = []
    # 段落
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            chunks.append(t)
    # 表格
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chunks.append(row_text)
    text = "\n".join(chunks)
    # 简单清理：连续空行/空格压缩
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text


def aggregate_major_text(major_name: str, main_path: Optional[str], sub_paths: list[str]) -> dict:
    """聚合一个专业的全部文本（主+特色班/方向班 取并集）。"""
    parts = []
    files_used = []

    if main_path:
        full = PLAN_ROOT / main_path
        text = parse_docx(full)
        if text:
            parts.append(f"【主培养方案】\n{text}")
            files_used.append(main_path)

    for sp in sub_paths:
        full = PLAN_ROOT / sp
        text = parse_docx(full)
        if text:
            tag = pathlib.Path(sp).stem
            parts.append(f"【{tag}】\n{text}")
            files_used.append(sp)

    return {
        "major_text": "\n\n========\n\n".join(parts),
        "files_used": files_used,
        "total_chars": sum(len(p) for p in parts),
    }


# ============================================================
# 3. 阶段 A：解析 + 聚合 → v4_supply_majors_raw.parquet
# ============================================================
def run_phase_a():
    print("=" * 60)
    print("阶段 A：解析 docx 并按专业聚合")
    print("=" * 60)

    rows = []
    for major, college, level, main, subs in MAJORS:
        agg = aggregate_major_text(major, main, subs)
        if not agg["files_used"]:
            print(f"  ✗ {major}（{college}）：无可用文件！")
            continue
        rows.append({
            "major_id": f"M_{len(rows):03d}",
            "专业名": major,
            "学院": college,
            "教育层次": level,
            "是否含子轨道": len(subs) > 0,
            "子轨道数": len(subs),
            "文件清单": json.dumps(agg["files_used"], ensure_ascii=False),
            "聚合文本": agg["major_text"],
            "字符数": agg["total_chars"],
        })
        print(f"  ✓ {major}（{level} / {college}）：{len(agg['files_used'])} 文件，{agg['total_chars']:,} 字符")

    df = pd.DataFrame(rows)
    df.to_parquet(OUT_RAW, index=False)
    print(f"\n阶段 A 完成：{len(df)} 个专业，输出 → {OUT_RAW}")
    print(f"  字符数：均值 {df['字符数'].mean():.0f}，中位 {df['字符数'].median():.0f}，最大 {df['字符数'].max():,}")
    return df


# ============================================================
# 4. 阶段 B：LLM 抽取（与需求侧 prompt 完全一致）
# ============================================================
SYSTEM_TOOLS = (
    "你是专业的高职专业建设分析师。"
    "请从人才培养方案中提取所有具体的技术工具、软件、框架、平台、编程语言名称（即课程会教学/学生应掌握的工具）。"
    "只返回JSON，不要任何额外说明。"
)
USER_TOOLS_TPL = (
    "从以下{count}份人才培养方案文本中，分别提取每份方案明确教学的具体技术工具/软件/编程语言列表。\n\n"
    "提取范围：编程语言（Python、Java、C++等）、软件（AutoCAD、SolidWorks、MATLAB等）、"
    "框架/库（Vue、React、TensorFlow等）、平台（AWS、阿里云、SAP等）、"
    "行业工具（PLC、ANSYS、Altium Designer、ERP等）。\n\n"
    "不要提取通用词（如\u300c计算机\u300d、\u300c系统\u300d、\u300c软件\u300d等），只要具体产品/工具名称。\n\n"
    "输出格式（JSON数组）：\n"
    '[{{"id": "记录ID", "tools": ["工具1", "工具2", ...]}}, ...]\n\n'
    "如果没有提及具体工具名称，输出空列表[]。\n\n"
    "人才培养方案文本：\n{texts}"
)

SYSTEM_SOFT = (
    "你是专业的高职专业建设分析师。"
    "请从人才培养方案中提取该专业拟培养的软技能/通用能力（即非专业技术能力，如沟通、团队协作、领导力、执行力等）。"
    "只返回JSON，不要任何额外说明。"
)
USER_SOFT_TPL = (
    "从以下{count}份人才培养方案中，分别提取每份方案明确强调的软技能/通用能力列表。\n\n"
    "软技能定义：非专业技术能力，包括但不限于：沟通能力、团队协作、执行力、学习能力、"
    "责任心、抗压能力、问题解决能力、领导力、客户服务意识、细心/严谨、创新思维等。\n\n"
    "输出格式（JSON数组）：\n"
    '[{{"id": "记录ID", "soft_skills": ["技能1", "技能2", ...]}}, ...]\n\n'
    "如果没有提及任何软技能，输出空列表[]。\n\n"
    "人才培养方案文本：\n{texts}"
)

DIGI_ATTRS = {
    "数字化程度": "该专业课程涉及数字化技术、信息系统、数字化工具的程度（0=完全不涉及，100=核心数字化专业）",
    "AI相关性": "该专业与人工智能、机器学习、数据科学的相关程度（0=完全不涉及，100=AI核心专业）",
    "技术复杂度": "该专业培养的技术技能深度和复杂程度（0=无技术要求，100=高度专业技术专业）",
}

# 与需求侧一致的归一化映射（从 09_soft_skills.py 与 12_tech_tools.py 提取）
SOFT_NORMALIZE = {
    "沟通协调能力": "沟通协调", "沟通能力": "沟通协调", "沟通协调": "沟通协调",
    "沟通指导能力": "沟通协调", "人际沟通": "沟通协调", "沟通表达": "沟通协调",
    "协调能力": "沟通协调",
    "团队协作能力": "团队协作", "团队合作": "团队协作", "协作精神": "团队协作",
    "团队精神": "团队协作",
    "学习能力": "学习能力", "自主学习": "学习能力", "持续学习": "学习能力",
    "学习意愿": "学习能力", "快速学习": "学习能力",
    "责任心": "责任心", "责任感": "责任心",
    "细心": "细心严谨", "严谨": "细心严谨", "细致": "细心严谨", "细心严谨": "细心严谨",
    "服务意识": "服务意识", "客户服务意识": "服务意识",
    "抗压能力": "抗压能力", "抱压": "抗压能力", "压力承受": "抗压能力", "压力管理": "抗压能力",
    "领导力": "领导力", "领导能力": "领导力", "管理能力": "领导力",
    "问题解决能力": "问题解决", "分析问题能力": "问题解决",
    "判断与决策能力": "判断决策", "决策能力": "判断决策",
    "创新能力": "创新思维", "创新思维": "创新思维",
    "执行力": "执行力", "计划与执行能力": "执行力",
}

TOOL_NORMALIZE_LOWER = {
    "python": "Python", "python3": "Python",
    "java": "Java", "c++": "C++", "c/c++": "C/C++", "c#": "C#",
    "javascript": "JavaScript", "js": "JavaScript",
    "typescript": "TypeScript", "ts": "TypeScript",
    "vue": "Vue.js", "vue.js": "Vue.js", "vue3": "Vue.js",
    "react": "React", "react.js": "React", "angular": "Angular",
    "tensorflow": "TensorFlow", "pytorch": "PyTorch",
    "autocad": "AutoCAD", "cad": "AutoCAD",
    "solidworks": "SolidWorks", "solid works": "SolidWorks",
    "matlab": "MATLAB", "proe": "Pro/E", "pro/e": "Pro/E",
    "catia": "CATIA", "ug": "UG/NX", "nx": "UG/NX",
    "ansys": "ANSYS",
    "altium": "Altium Designer", "altium designer": "Altium Designer",
    "sap": "SAP", "plc": "PLC",
    "mysql": "MySQL", "postgresql": "PostgreSQL", "redis": "Redis", "mongodb": "MongoDB",
    "docker": "Docker", "kubernetes": "Kubernetes", "k8s": "Kubernetes",
    "git": "Git", "linux": "Linux",
    "android": "Android", "ios": "iOS",
    "office": "Microsoft Office", "excel": "Excel", "word": "Word", "powerpoint": "PowerPoint",
    "erp": "ERP", "金蝶": "金蝶", "用友": "用友",
    "photoshop": "Photoshop", "ps": "Photoshop",
    "illustrator": "Illustrator", "ai": "Illustrator",
    "premiere": "Premiere", "pr": "Premiere",
    "after effects": "After Effects", "ae": "After Effects",
    "3dmax": "3ds Max", "3ds max": "3ds Max", "3dsmax": "3ds Max",
    "maya": "Maya", "blender": "Blender", "unity": "Unity", "unreal": "Unreal Engine",
}


def normalize_tools(items: list) -> list:
    out = []
    for x in items or []:
        if not isinstance(x, str):
            continue
        key = x.strip().lower()
        out.append(TOOL_NORMALIZE_LOWER.get(key, x.strip()))
    # 去重，保持顺序
    seen, dedup = set(), []
    for x in out:
        if x not in seen:
            seen.add(x)
            dedup.append(x)
    return dedup


def normalize_soft(items: list) -> list:
    out = []
    for x in items or []:
        if not isinstance(x, str):
            continue
        out.append(SOFT_NORMALIZE.get(x.strip(), x.strip()))
    seen, dedup = set(), []
    for x in out:
        if x not in seen:
            seen.add(x)
            dedup.append(x)
    return dedup


# ============================================================
# 5. 阶段 B 主流程
# ============================================================
def run_phase_b():
    print("\n" + "=" * 60)
    print("阶段 B：LLM 抽取（工具 / 软技能 / 数字化评分）")
    print("=" * 60)

    df = pd.read_parquet(OUT_RAW)
    df = df.copy()
    # 培养方案文本通常很长。我们按"分批 × 单条上限"控制 token：
    # 单条最多 5000 字符（约 2.5K tokens）× batch_size=3 = 单批 ~7.5K tokens 输入，安全。
    df["分析文本"] = df["聚合文本"].str.slice(0, 5000)

    # 1) 工具抽取
    print("\n[B-1] 工具抽取 …")
    tools_df = batch_extract(
        df=df,
        id_column="major_id",
        text_column="分析文本",
        task_name="v4_supply_tools",
        system_prompt=SYSTEM_TOOLS,
        user_prompt_template=USER_TOOLS_TPL,
        output_field="tools",
        batch_size=3,
        text_truncate=5000,
        reset=True,
    )

    # 2) 软技能抽取
    print("\n[B-2] 软技能抽取 …")
    soft_df = batch_extract(
        df=df,
        id_column="major_id",
        text_column="分析文本",
        task_name="v4_supply_soft",
        system_prompt=SYSTEM_SOFT,
        user_prompt_template=USER_SOFT_TPL,
        output_field="soft_skills",
        batch_size=3,
        text_truncate=5000,
        reset=True,
    )

    # 3) 数字化评分
    print("\n[B-3] 数字化评分 …")
    digi_df = batch_rate(
        df=df,
        id_column="major_id",
        text_column="分析文本",
        task_name="v4_supply_digi",
        attributes=DIGI_ATTRS,
        batch_size=3,
        text_truncate=5000,
        reset=True,
    )

    # 合并
    out = (
        df[["major_id", "专业名", "学院", "教育层次", "是否含子轨道", "子轨道数", "文件清单", "字符数"]]
        .merge(tools_df, on="major_id", how="left")
        .merge(soft_df, on="major_id", how="left")
        .merge(digi_df, on="major_id", how="left")
    )

    # 归一化（强制为 Python list，避免 pyarrow 混合类型异常）
    out["tools"] = out["tools"].apply(lambda x: list(x) if hasattr(x, "__iter__") and not isinstance(x, str) else [])
    out["soft_skills"] = out["soft_skills"].apply(lambda x: list(x) if hasattr(x, "__iter__") and not isinstance(x, str) else [])
    out["tools_normalized"] = out["tools"].apply(normalize_tools)
    out["soft_skills_normalized"] = out["soft_skills"].apply(normalize_soft)

    out.to_parquet(OUT_FINAL, index=False)
    print(f"\n阶段 B 完成：输出 → {OUT_FINAL}")
    print(f"  样本：{len(out)} 个专业；含工具列表者 {out['tools_normalized'].apply(bool).sum()}；")
    print(f"        含软技能列表者 {out['soft_skills_normalized'].apply(bool).sum()}")
    return out


# ============================================================
# 6. 入口
# ============================================================
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--phase", choices=["a", "b", "all"], default="all",
                        help="a=仅解析 docx; b=仅 LLM 抽取（需先有 raw）; all=完整流程")
    args = parser.parse_args()

    if args.phase in ("a", "all"):
        run_phase_a()
    if args.phase in ("b", "all"):
        run_phase_b()

    print("\n=== 完成 ===")
